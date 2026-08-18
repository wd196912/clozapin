#!/usr/bin/env python3
"""
Regenerate Supplementary Table S1 (clozapine vs olanzapine) as docx.
Reads disproportionality_clo_vs_ola_ascii.json (primaryid-level denominator,
13,691 olanzapine primary suspect reports — consistent with the report-level
clozapine 44,055 / risperidone 15,130 denominators).
New signal rule: ROR/PRR reported only when olanzapine events >= 5;
positive = IC025 > 0 AND olanzapine events >= 5.
"""

import json
import os

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_DIR = "F:/clo-ola/data"
OUT_PATH = "F:/clazpin/manuscript/analysis/tables/table_s1_olanzapine_sensitivity.docx"

ORDER = [
    "Composite pulmonary infection endpoint (primary)",
    "Pneumonia spectrum (subgroup)",
    "Non-pneumonia lower respiratory tract infection (subgroup)",
    "Special/other pulmonary infection terms (subgroup)",
    "respiratory tract infection", "pneumonitis", "empyema",
    "pneumonia klebsiella", "pneumonia viral",
    "respiratory tract infection viral", "pneumonia influenzal",
    "pneumonia staphylococcal", "pulmonary tuberculosis",
    "idiopathic interstitial pneumonia", "lung abscess",
    "lower respiratory tract infection", "upper respiratory tract infection",
    "covid-19 pneumonia", "pneumonia", "pneumonia aspiration",
    "pneumonia bacterial",
    "Other pulmonary infection terms (9 additional PTs)",
]


def title_case(label):
    if label.startswith(("Composite", "Pneumonia spectrum", "Non-pneumonia",
                         "Special/other", "Other pulmonary")):
        return label
    return label.replace("-", " ").title().replace("Covid", "COVID")


def main():
    with open(os.path.join(DATA_DIR, "disproportionality_clo_vs_ola_ascii.json"),
              encoding="utf-8") as f:
        data = json.load(f)
    results = {r["label"]: r for r in data["primaryid_level"]["results"]}
    n_ola = data["primaryid_level"]["olanzapine_total"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    p_title = doc.add_paragraph()
    run = p_title.add_run(
        "Supplementary Table S1. Sensitivity Analysis: Clozapine vs Olanzapine "
        "Disproportionality Analysis (FAERS ASCII, 2022–2025)")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"

    headers = ["Preferred Term / Endpoint", "Clozapine\n(n)", "Olanzapine\n(n)",
               "ROR", "95% CI", "PRR", "IC", "IC₀₂₅", "Positive\nSignal"]

    rows = []
    for label in ORDER:
        d = results[label]
        ror_s = f"{d['ror']:.2f}" if d["ror"] != "NR" else "NR"
        ci_s = (f"{d['ror_ci_95_low']:.2f}–{d['ror_ci_95_high']:.2f}"
                if d["ror_ci_95_low"] != "NR" else "—")
        prr_s = f"{d['prr']:.2f}" if d["prr"] != "NR" else "NR"
        sig_s = "Yes" if d["signal_positive"] else "No"
        rows.append([title_case(label), str(d["a"]), str(d["c"]), ror_s, ci_s,
                     prr_s, f"{d['ic']:.3f}", f"{d['ic025']:.3f}", sig_s])

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(7)
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shading)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(7)
            run.font.name = "Arial"
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if j > 0
                           else WD_ALIGN_PARAGRAPH.LEFT)
            if j == 8 and val == "Yes":
                run.bold = True
        if i % 2 == 0:
            for j in range(len(headers)):
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F2F2")
                table.rows[i + 1].cells[j]._tc.get_or_add_tcPr().append(shading)

    p_fn = doc.add_paragraph()
    run = p_fn.add_run(
        "ROR = Reporting Odds Ratio; CI = Confidence Interval; PRR = Proportional "
        "Reporting Ratio; IC = Information Component; IC₀₂₅ = lower bound of 95% "
        "credibility interval. FAERS ASCII data (44,055 clozapine primary suspect "
        f"reports; {n_ola:,} olanzapine primary suspect reports, 2022Q1–2025Q4). "
        "ROR/PRR are not reported (NR) when the olanzapine group has fewer than 5 "
        "events; a positive signal is defined as IC₀₂₅ > 0 with ≥5 comparator "
        "events. The composite endpoint comprised pneumonia, pneumonia aspiration, "
        "lower respiratory tract infection, pneumonia bacterial, pneumonia viral, "
        "empyema, lung abscess, pneumonia klebsiella, pneumonia influenzal, and "
        "pneumonia staphylococcal.")
    run.font.size = Pt(7)
    run.font.name = "Arial"
    run.italic = True

    doc.save(OUT_PATH)
    print(f"Saved S1 ({len(rows)} data rows, olanzapine N={n_ola}) to {OUT_PATH}")


if __name__ == "__main__":
    main()
