#!/usr/bin/env python3
"""
Generate Supplementary Table S4 (classification-scheme sensitivity) as docx.
Three schemes: A original per-PT (old rule), B composite endpoint,
C clinical subgroups (new rule).
All numbers ASCII-derived from recompute_composite_endpoint.R and the
pre-composite Table 2.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "F:/clazpin/manuscript/analysis/tables/table_s4_classification_sensitivity.docx"

ROWS = [
    # scheme, event definition, cloz, risp, ROR (95% CI), IC (IC025), positive
    ("A — Original per-PT", "Lower respiratory tract infection", "237", "1",
     "81.83 (11.48–583.38)", "0.419 (0.138)", "Yes"),
    ("A — Original per-PT", "Pneumonia", "723", "75",
     "3.35 (2.64–4.25)", "0.283 (0.127)", "Yes"),
    ("A — Original per-PT", "Pneumonia aspiration", "209", "40",
     "1.80 (1.28–2.52)", "0.173 (−0.112)", "No"),
    ("A — Original per-PT", "Upper respiratory tract infection", "27", "2",
     "4.64 (1.10–19.51)", "0.316 (−0.492)", "No"),
    ("A — Original per-PT", "COVID-19 pneumonia", "25", "2",
     "4.29 (1.02–18.13)", "0.308 (−0.530)", "No"),
    ("A — Original per-PT", "Pneumonia bacterial", "12", "3",
     "1.37 (0.39–4.87)", "0.100 (−1.051)", "No"),
    ("A — Original per-PT", "11 remaining PTs (0 risperidone events each)",
     "2–53", "0", "Not estimable", "0.330–0.421 (all IC₀₂₅ < 0)", "No"),
    ("B — Composite endpoint", "Composite pulmonary infection endpoint (10 PTs)",
     "1177", "115", "3.58 (2.96–4.34)", "0.291 (0.168)", "Yes"),
    ("C — Clinical subgroups", "Pneumonia spectrum (9 PTs)", "963", "114",
     "2.94 (2.42–3.58)", "0.264 (0.129)", "Yes"),
    ("C — Clinical subgroups", "Non-pneumonia lower respiratory tract "
     "infection (1 PT)", "237", "1", "NR (<5 comparator events)",
     "0.419 (0.138)", "No"),
    ("C — Clinical subgroups", "Special/other pulmonary infection terms (16 PTs)",
     "145", "4", "NR (<5 comparator events)", "0.385 (0.029)", "No"),
]


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    p_title = doc.add_paragraph()
    run = p_title.add_run(
        "Supplementary Table S4. Sensitivity of Disproportionality Results to "
        "the Classification Scheme")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"

    headers = ["Classification scheme", "Event definition", "Clozapine\n(n)",
               "Risperidone\n(n)", "ROR (95% CI)", "IC (IC₀₂₅)",
               "Positive\nsignal"]

    table = doc.add_table(rows=len(ROWS) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(7)
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shading)

    for i, row_data in enumerate(ROWS):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(7)
            run.font.name = "Arial"
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if j >= 2
                           else WD_ALIGN_PARAGRAPH.LEFT)
        if i % 2 == 0:
            for j in range(len(headers)):
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F2F2")
                table.rows[i + 1].cells[j]._tc.get_or_add_tcPr().append(shading)

    p_fn = doc.add_paragraph()
    run = p_fn.add_run(
        "ROR = Reporting Odds Ratio; CI = Confidence Interval; IC = Information "
        "Component; IC₀₂₅ = lower bound of 95% credibility interval; NR = not "
        "reported; PT = MedDRA Preferred Term. Scheme A applied the original "
        "per-PT analysis with the original positivity criteria (ROR > 1 with "
        "lower 95% CI > 1 and IC₀₂₅ > 0); Schemes B and C applied the revised "
        "positivity criteria (IC₀₂₅ > 0 with ≥5 comparator events), under which "
        "ROR/PRR are not reported when the comparator group has fewer than 5 "
        "events. Pneumonia showed a positive signal under all three schemes, and "
        "the composite endpoint and pneumonia spectrum subgroup were positive "
        "under the revised criteria; the lower respiratory tract infection "
        "estimate was unstable under Scheme A (ROR 81.83, based on a single "
        "risperidone event) and was not classified as a signal under the revised "
        "rule. Clozapine N = 44,055; risperidone N = 15,130.")
    run.font.size = Pt(7)
    run.font.name = "Arial"
    run.italic = True

    doc.save(OUT)
    print(f"Saved S4 ({len(ROWS)} data rows) to {OUT}")


if __name__ == "__main__":
    main()
