#!/usr/bin/env python3
"""
Rebuild Tables 1 and 2 in tables_1_3.docx for the composite-endpoint revision.
Table 1: infection-type rows replaced by principal-diagnosis counts (sums to 1,305).
Table 2: composite endpoint + subgroups added above the 17 PTs; new NR rule
(comparator n<5 -> ROR/PRR not reported); vs-OLA columns from primaryid-level
olanzapine denominator (13,691). Table 3 copied unchanged.
Inputs:
  F:/clazpin/composite_endpoint_signal_results.csv   (vs risperidone, new rule)
  F:/clo-ola/data/disproportionality_clo_vs_ola_ascii.json (vs olanzapine)
"""

import csv
import os
import json
import shutil

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = "F:/clazpin/manuscript/analysis/tables/tables_1_3.docx"
BACKUP = "F:/clazpin/manuscript/analysis/tables/tables_1_3_precomposite_backup.docx"
CSV_RISP = "F:/clazpin/composite_endpoint_signal_results.csv"
JSON_OLA = "F:/clo-ola/data/disproportionality_clo_vs_ola_ascii.json"

# Principal-diagnosis Table 1 rows: label, n, fatal, nonfatal
T1_INFECTION = [
    ("Infection: Pneumonia", 746, 166, 580),
    ("Infection: Pneumonia aspiration", 209, 93, 116),
    ("Infection: Lower respiratory tract infection", 214, 34, 180),
    ("Infection: COVID-19 pneumonia", 21, 10, 11),
    ("Infection: Respiratory tract infection", 51, 4, 47),
    ("Infection: Upper respiratory tract infection", 22, 2, 20),
    ("Infection: Pneumonitis", 17, 0, 17),
    ("Infection: Empyema", 14, 0, 14),
    ("Infection: Pulmonary tuberculosis", 6, 0, 6),
    ("Infection: Lung abscess", 3, 0, 3),
    ("Infection: Pulmonary sepsis", 2, 0, 2),
    ("Infection: Idiopathic interstitial pneumonia", 0, 0, 0),
]

N_ALL, N_FATAL, N_NONFATAL = 1305, 309, 996

# Table 2 row order: 4 endpoint rows + 17 PTs (old table order) + 9-additional row
T2_ORDER = [
    "Composite pulmonary infection endpoint (primary)",
    "Pneumonia spectrum (subgroup)",
    "Non-pneumonia lower respiratory tract infection (subgroup)",
    "Special/other pulmonary infection terms (subgroup)",
    "respiratory tract infection", "pneumonitis", "empyema",
    "pneumonia klebsiella", "pneumonia viral",
    "respiratory tract infection viral", "pneumonia influenzal",
    "pneumonia staphylococcal", "pulmonary tuberculosis",
    "idiopathic interstitial pneumonia", "lung abscess",
    "lower respiratory tract infection", "upper respiratory tract infection",
    "covid-19 pneumonia", "pneumonia", "pneumonia aspiration",
    "pneumonia bacterial",
    "Other pulmonary infection terms (9 additional PTs)",
]

T2_DISPLAY = {
    "Composite pulmonary infection endpoint (primary)":
        "Composite Pulmonary Infection Endpoint (Primary)",
    "Pneumonia spectrum (subgroup)": "Pneumonia Spectrum (Subgroup)",
    "Non-pneumonia lower respiratory tract infection (subgroup)":
        "Non-Pneumonia Lower Respiratory Tract Infection (Subgroup)",
    "Special/other pulmonary infection terms (subgroup)":
        "Special/Other Pulmonary Infection Terms (Subgroup)",
    "Other pulmonary infection terms (9 additional PTs)":
        "Other Pulmonary Infection Terms (9 Additional PTs)",
    "covid-19 pneumonia": "Covid-19 Pneumonia",
}

T2_TITLE_NEW = ("Table 2. Disproportionality Analysis of the Composite Pulmonary "
                "Infection Endpoint, Clinical Subgroups, and Individual Preferred "
                "Terms: Clozapine vs. Risperidone (Primary) and Clozapine vs. "
                "Olanzapine (Sensitivity)")

T2_FOOTNOTE_NEW = ("ROR = Reporting Odds Ratio; CI = Confidence Interval; "
                   "IC = Information Component; IC₀₂₅ = lower bound of 95% "
                   "credibility interval. † ROR and PRR are not reported when the "
                   "comparator group has fewer than 5 events; the IC is reported "
                   "alone. A positive signal is defined as IC₀₂₅ > 0 with "
                   "≥5 comparator events. The composite endpoint (primary "
                   "analysis) comprised pneumonia, pneumonia aspiration, lower "
                   "respiratory tract infection, pneumonia bacterial, pneumonia "
                   "viral, empyema, lung abscess, pneumonia klebsiella, pneumonia "
                   "influenzal, and pneumonia staphylococcal; the pneumonia "
                   "spectrum subgroup excluded lower respiratory tract infection; "
                   "the special/other subgroup comprised the remaining 16 "
                   "screening PTs. Individual PT analyses were exploratory. "
                   "Clozapine N = 44,055; risperidone N = 15,130; olanzapine "
                   "N = 13,691 (report level).")

T1_NOTE_NEW = ("Infection types reflect the principal diagnosis per report, "
               "assigned by a predefined priority chain (pneumonia aspiration > "
               "pneumonia > lower respiratory tract infection > COVID-19 "
               "pneumonia > upper respiratory tract infection > respiratory "
               "tract infection > pneumonitis > empyema > lung abscess > "
               "pulmonary tuberculosis > pulmonary sepsis > idiopathic "
               "interstitial pneumonia); counts therefore sum exactly to 1,305.")


def pct(x, den):
    return f"{100 * x / den:.1f}%"


def set_cell(cell, text, font_name, font_size, align_center, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold


def remove_row(table, index):
    tr = table.rows[index]._tr
    table._tbl.remove(tr)


def append_row(table, values, font_name, font_size, center_from=0):
    row = table.add_row()
    for j, val in enumerate(values):
        set_cell(row.cells[j], str(val), font_name, font_size,
                 align_center=(j >= center_from))
    return row


def get_cell_font(table, row_idx, col_idx):
    cell = table.rows[row_idx].cells[col_idx]
    for p in cell.paragraphs:
        for r in p.runs:
            if r.text.strip():
                return r.font.name or "Arial", r.font.size or Pt(7)
    return "Arial", Pt(7)


def main():
    if not os.path.exists(BACKUP):
        shutil.copy(DOCX, BACKUP)
        print(f"Backed up existing docx to {BACKUP}")
    else:
        shutil.copy(BACKUP, DOCX)
        print(f"Restored {DOCX} from backup")

    with open(CSV_RISP, "r", encoding="utf-8-sig") as f:
        risp = {r["Label"]: r for r in csv.DictReader(f)}
    with open(JSON_OLA, "r", encoding="utf-8") as f:
        ola_data = json.load(f)
    ola = {r["label"]: r for r in ola_data["primaryid_level"]["results"]}

    doc = Document(DOCX)
    t1, t2, t3 = doc.tables[0], doc.tables[1], doc.tables[2]

    # ---------- Table 1: replace infection rows ----------
    font1, size1 = get_cell_font(t1, 1, 0)
    # old infection rows are the last 5 data rows (indices 19-23, 0-based)
    for _ in range(5):
        remove_row(t1, 19)
    for label, n, fatal, nonfatal in T1_INFECTION:
        values = [
            label,
            f"{n} ({pct(n, N_ALL)})",
            f"{fatal} ({pct(fatal, N_FATAL)})" if fatal > 0 else "0 (0.0%)",
            f"{nonfatal} ({pct(nonfatal, N_NONFATAL)})" if nonfatal > 0 else "0 (0.0%)",
            "—",
        ]
        append_row(t1, values, font1, size1, center_from=1)

    # add principal-diagnosis note paragraph after the weight footnote
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(0)
    # move it right after the existing "† Weight data" paragraph
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("† Weight data"):
            p._p.addnext(note_p._p)
            break
    run = note_p.add_run(T1_NOTE_NEW)
    run.font.name = font1
    run.font.size = size1
    run.italic = True

    # ---------- Table 2: title, rows, footnote ----------
    for p in doc.paragraphs:
        if p.text.startswith("Table 2. Disproportionality Analysis"):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = T2_TITLE_NEW
                break

    font2, size2 = get_cell_font(t2, 1, 0)
    while len(t2.rows) > 1:
        remove_row(t2, 1)
    for i, label in enumerate(T2_ORDER):
        r_risp = risp[label]
        r_ola = ola[label]
        disp = T2_DISPLAY.get(label, label.replace("-", " ").title())

        def risp_ror():
            if r_risp["ROR"] == "NR":
                return "NR†", "—"
            return (f"{float(r_risp['ROR']):.2f}",
                    f"{float(r_risp['ROR_CI95_low']):.2f}–"
                    f"{float(r_risp['ROR_CI95_high']):.2f}")

        ror_s, ci_s = risp_ror()
        sig_risp = "Positive" if r_risp["Signal"] == "Positive" else "—"
        ola_ror = r_ola["ror"]
        ola_ci = r_ola["ror_ci_95_low"]
        if ola_ror != "NR":
            ola_ror_s = f"{ola_ror:.2f}"
            ola_ci_s = f"{ola_ci:.2f}–{r_ola['ror_ci_95_high']:.2f}"
        else:
            ola_ror_s, ola_ci_s = "NR†", "—"
        sig_ola = "Positive" if r_ola["signal_positive"] else "—"

        values = [
            disp, r_risp["a"], r_risp["c"], ror_s, ci_s,
            f"{float(r_risp['IC']):.3f}", f"{float(r_risp['IC_025']):.3f}",
            sig_risp,
            ola_ror_s, ola_ci_s, f"{r_ola['ic025']:.3f}", sig_ola,
        ]
        row = append_row(t2, values, font2, size2, center_from=1)
        if label.startswith("Composite"):
            set_cell(row.cells[0], disp, font2, size2,
                     align_center=False, bold=True)

    for p in doc.paragraphs:
        if p.text.strip().startswith("ROR = Reporting Odds Ratio"):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = T2_FOOTNOTE_NEW
                break

    doc.save(DOCX)
    print(f"Saved {DOCX}: Table 1 = {len(t1.rows)} rows, "
          f"Table 2 = {len(t2.rows)} rows, Table 3 = {len(t3.rows)} rows (unchanged)")


if __name__ == "__main__":
    main()
