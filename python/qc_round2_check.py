#!/usr/bin/env python3
"""Round-2 QC: recompute all new numbers from source counts and compare with rendered docx."""
import csv, json, math, sys

from docx import Document

DOCX = 'F:/clazpin/manuscript/final_comp.docx'
T1_CSV = 'F:/clazpin/principal_diagnosis_table1.csv'
SIG_CSV = 'F:/clazpin/composite_endpoint_signal_results.csv'
OLA_JSON = 'F:/clo-ola/data/disproportionality_clo_vs_ola_ascii.json'

N_ALL, N_FATAL, N_NONFATAL = 1305, 309, 996
N_CLO, N_RISP = 44055, 15130
N_OLA = 13691

fails = []

def check(name, cond, detail=''):
    status = 'OK ' if cond else 'FAIL'
    print(f'[{status}] {name} {detail}')
    if not cond:
        fails.append(name)

# ---------- 1) Table 1 principal-diagnosis sums & percentages ----------
rows = list(csv.DictReader(open(T1_CSV, encoding='utf-8-sig')))
tot_n = sum(int(r['n']) for r in rows)
tot_f = sum(int(r['fatal']) for r in rows)
check('T1 sum n = 1305', tot_n == N_ALL, f'({tot_n})')
check('T1 sum fatal = 309', tot_f == N_FATAL, f'({tot_f})')

doc = Document(DOCX)
doc_text = []
for p in doc.paragraphs:
    doc_text.append(p.text)
for t in doc.tables:
    for row in t.rows:
        doc_text.append(' | '.join(c.text for c in row.cells))
doc_txt = '\n'.join(doc_text)

def pct_str(x, den):
    return f'{100 * x / den:.1f}%'

t1_ok = True
for r in rows:
    n, f = int(r['n']), int(r['fatal'])
    nf = n - f
    expected = [f'{n} ({pct_str(n, N_ALL)})', f'{f} ({pct_str(f, N_FATAL)})',
                f'{nf} ({pct_str(nf, N_NONFATAL)})']
    label = r['principal']
    for line in doc_txt.splitlines():
        if line.startswith('Infection: ' + label + ' |'):
            cells = line.split(' | ')
            for i, exp in enumerate(expected, start=1):
                if cells[i] != exp:
                    check(f'T1 pct {label} col{i}', False,
                          f'rendered={cells[i]} expected={exp}')
                    t1_ok = False
            break
check('T1 all percentages correct', t1_ok)

# ---------- 2) Table 2: ROR back-solve from 2x2 ----------
sig = {r['Label']: r for r in csv.DictReader(open(SIG_CSV, encoding='utf-8-sig'))}
for label, r in sig.items():
    a, c = int(r['a']), int(r['c'])
    b, d = N_CLO - a, N_RISP - c
    if r['ROR'] == 'NR':
        continue
    ror = (a * d) / (b * c)
    shown = float(r['ROR'])
    ok = abs(ror - shown) < 0.011
    check(f'T2 ROR backsolve {label[:45]}', ok,
          f'calc={ror:.3f} shown={shown}')
    # CI check via SE
    se = math.sqrt(1/a + 1/b + 1/c + 1/d)
    lo = math.exp(math.log(ror) - 1.96 * se)
    hi = math.exp(math.log(ror) + 1.96 * se)
    ok2 = abs(lo - float(r['ROR_CI95_low'])) < 0.011 and abs(hi - float(r['ROR_CI95_high'])) < 0.011
    check(f'T2 CI {label[:45]}', ok2, f'calc=({lo:.2f}-{hi:.2f}) shown=({r["ROR_CI95_low"]}-{r["ROR_CI95_high"]})')

# ---------- 3) S1 (vs olanzapine): ROR back-solve + rule ----------
ola = json.load(open(OLA_JSON, encoding='utf-8'))['primaryid_level']['results']
for x in ola:
    a, c = int(x['a']), int(x['c'])
    b, d = N_CLO - a, N_OLA - c
    if x['ror'] == 'NR':
        assert c < 5, f'NR but c={c} for {x["label"]}'
        continue
    ror = (a * d) / (b * c)
    ok = abs(ror - float(x['ror'])) < 0.011
    check(f'S1 ROR backsolve {x["label"][:45]}', ok,
          f'calc={ror:.3f} shown={x["ror"]}')
    ok_rule = c >= 5
    check(f'S1 reported iff c>=5 {x["label"][:40]}', ok_rule, f'(c={c})')

# ---------- 4) text percentages in manuscript paragraphs ----------
txt_checks = [
    ('746 (57.2%)', pct_str(746, N_ALL)),
    ('214 (16.4%)', pct_str(214, N_ALL)),
    ('209 (16.0%)', pct_str(209, N_ALL)),
    ('10/21, 47.6%', pct_str(10, 21)),
    ('93/209, 44.5%', pct_str(93, 209)),
    ('166/746, 22.3%', pct_str(166, 746)),
    ('34/214, 15.9%', pct_str(34, 214)),
]
for needle, expected in txt_checks:
    ok = needle in doc_txt
    check(f'text "{needle}"', ok)
    if ok:
        check(f'text "{needle}" pct correct', expected in needle)

# composite endpoint counts in text (normalize NBSP that pandoc inserts after "vs.")
norm = doc_txt.replace(' ', ' ')
for needle in ['1,177 clozapine vs. 115 risperidone', '963 vs. 114 events',
               '237 vs. 1 event', '145 vs. 4 events', '723 vs. 75 events',
               '209 vs. 40 events', '19 olanzapine events',
               'ROR 3.89, 95% CI 2.44–6.21; IC₀₂₅ 0.006']:
    check(f'text "{needle}"', needle in norm)

print()
print('FAILURES:', len(fails))
if fails:
    print('\n'.join(fails))
    sys.exit(1)
